from __future__ import annotations

import hashlib
import json
import re
from pathlib import Path

import fitz
import pandas as pd
from docx import Document

from .arabic import normalize_arabic_text
from .io_utils import read_text

SUPPORTED_SUFFIXES = {".txt", ".docx", ".pdf"}


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


def extract_pdf_text(path: str | Path) -> str:
    document = fitz.open(path)
    try:
        return "\n\n".join(page.get_text("text") for page in document)
    finally:
        document.close()


def extract_docx_text(path: str | Path) -> str:
    document = Document(path)
    blocks = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values:
                blocks.append("\t".join(values))
    return "\n\n".join(blocks)


def extract_text(path: str | Path) -> str:
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".txt":
        return read_text(path)
    if suffix == ".docx":
        return extract_docx_text(path)
    if suffix == ".pdf":
        return extract_pdf_text(path)
    raise ValueError(f"Unsupported source file: {path}")


def normalize_document_text(text: str, language: str) -> str:
    text = text.replace("\x00", " ").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    if language == "Arabic":
        text = normalize_arabic_text(text)
    return text


def remove_common_non_authorial_blocks(text: str) -> str:
    lines = text.splitlines()
    output: list[str] = []
    skipping_references = False
    reference_headers = {"references", "bibliography", "المراجع", "المصادر"}
    for line in lines:
        stripped = line.strip()
        lowered = stripped.lower().rstrip(":")
        if lowered in reference_headers:
            skipping_references = True
            continue
        if skipping_references:
            continue
        if re.fullmatch(r"(?:table of contents|contents|فهرس المحتويات)", lowered):
            continue
        if re.fullmatch(r"(?:page\s+)?\d+", lowered):
            continue
        output.append(line)
    return re.sub(r"\n{3,}", "\n\n", "\n".join(output)).strip()


def build_manifest(root: str | Path) -> pd.DataFrame:
    root = Path(root)
    rows: list[dict[str, object]] = []
    for index, path in enumerate(sorted(p for p in root.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES), start=1):
        relative = path.relative_to(root)
        parts = relative.parts
        inferred_language = "Arabic" if any("arab" in part.lower() or "عرب" in part for part in parts) else ("Code" if any("code" in part.lower() or "coding" in part.lower() for part in parts) else "English")
        rows.append(
            {
                "file_id": f"P2-{index:04d}",
                "file_name": path.name,
                "file_path": relative.as_posix(),
                "language": inferred_language,
                "university": parts[0] if len(parts) > 1 else "",
                "major": "",
                "year_grouping": "",
                "assignment_type": "",
                "source_format": path.suffix.lower().lstrip("."),
            }
        )
    return pd.DataFrame(rows)


def deidentified_id(relative_path: str) -> str:
    return hashlib.sha256(relative_path.encode("utf-8")).hexdigest()[:16]


def preprocess_manifest(
    manifest: pd.DataFrame,
    *,
    files_root: str | Path,
    output_root: str | Path,
    minimum_words: int = 300,
    remove_non_authorial: bool = True,
) -> pd.DataFrame:
    files_root = Path(files_root)
    output_root = Path(output_root)
    output_root.mkdir(parents=True, exist_ok=True)
    rows: list[dict[str, object]] = []
    for item in manifest.to_dict(orient="records"):
        source = Path(item["file_path"])
        if not source.is_absolute():
            source = files_root / source
        record = dict(item)
        try:
            text = extract_text(source)
            text = normalize_document_text(text, str(item.get("language", "English")))
            if remove_non_authorial:
                text = remove_common_non_authorial_blocks(text)
            words = count_words(text)
            output_name = f"{item.get('file_id') or deidentified_id(str(item['file_path']))}.txt"
            destination = output_root / output_name
            if words >= minimum_words:
                destination.write_text(text, encoding="utf-8")
                record.update({"processed_path": destination.as_posix(), "word_count": words, "preprocess_status": "included", "preprocess_error": ""})
            else:
                record.update({"processed_path": "", "word_count": words, "preprocess_status": "excluded_below_minimum_words", "preprocess_error": ""})
        except Exception as error:
            record.update({"processed_path": "", "word_count": 0, "preprocess_status": "error", "preprocess_error": repr(error)})
        rows.append(record)
    return pd.DataFrame(rows)


def split_by_paragraphs(text: str, max_characters: int = 6000) -> list[str]:
    paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", text) if paragraph.strip()]
    chunks: list[str] = []
    current: list[str] = []
    current_length = 0
    for paragraph in paragraphs:
        required = len(paragraph) + (2 if current else 0)
        if current and current_length + required > max_characters:
            chunks.append("\n\n".join(current))
            current, current_length = [], 0
        if len(paragraph) > max_characters:
            if current:
                chunks.append("\n\n".join(current))
                current, current_length = [], 0
            for start in range(0, len(paragraph), max_characters):
                chunks.append(paragraph[start : start + max_characters])
        else:
            current.append(paragraph)
            current_length += required
    if current:
        chunks.append("\n\n".join(current))
    return chunks
